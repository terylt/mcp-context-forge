#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Location: ./mcp-servers/python/docx_server/src/docx_server/server.py
Copyright 2025
SPDX-License-Identifier: Apache-2.0
Authors: Mihai Criveti

DOCX MCP Server

A comprehensive MCP server for creating, editing, and analyzing Microsoft Word (.docx) documents.
Provides tools for document creation, text manipulation, formatting, and document analysis.
"""

import asyncio
import json
import logging
import sys
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from mcp.server import Server
from mcp.server.models import InitializationOptions
from mcp.types import EmbeddedResource, ImageContent, TextContent, Tool
from pydantic import BaseModel, Field

# Configure logging to stderr to avoid MCP protocol interference
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stderr)],
)
logger = logging.getLogger(__name__)

# Create server instance
server = Server("docx-server")


class DocumentRequest(BaseModel):
    """Base request for document operations."""
    file_path: str = Field(..., description="Path to the DOCX file")


class CreateDocumentRequest(DocumentRequest):
    """Request to create a new document."""
    title: str | None = Field(None, description="Document title")
    author: str | None = Field(None, description="Document author")


class AddTextRequest(DocumentRequest):
    """Request to add text to a document."""
    text: str = Field(..., description="Text to add")
    paragraph_index: int | None = Field(None, description="Paragraph index to insert at (None for end)")
    style: str | None = Field(None, description="Style to apply")


class AddHeadingRequest(DocumentRequest):
    """Request to add a heading to a document."""
    text: str = Field(..., description="Heading text")
    level: int = Field(1, description="Heading level (1-9)", ge=1, le=9)


class FormatTextRequest(DocumentRequest):
    """Request to format text in a document."""
    paragraph_index: int = Field(..., description="Paragraph index to format")
    run_index: int | None = Field(None, description="Run index within paragraph (None for entire paragraph)")
    bold: bool | None = Field(None, description="Make text bold")
    italic: bool | None = Field(None, description="Make text italic")
    underline: bool | None = Field(None, description="Underline text")
    font_size: int | None = Field(None, description="Font size in points")
    font_name: str | None = Field(None, description="Font name")


class AddTableRequest(DocumentRequest):
    """Request to add a table to a document."""
    rows: int = Field(..., description="Number of rows", ge=1)
    cols: int = Field(..., description="Number of columns", ge=1)
    data: list[list[str]] | None = Field(None, description="Table data (optional)")
    headers: list[str] | None = Field(None, description="Column headers (optional)")


class AnalyzeDocumentRequest(DocumentRequest):
    """Request to analyze document content."""
    include_structure: bool = Field(True, description="Include document structure analysis")
    include_formatting: bool = Field(True, description="Include formatting analysis")
    include_statistics: bool = Field(True, description="Include text statistics")


class DocumentOperation:
    """Handles document operations."""

    @staticmethod
    def create_document(file_path: str, title: str | None = None, author: str | None = None) -> dict[str, Any]:
        """Create a new DOCX document."""
        try:
            # Create document
            doc = Document()

            # Set document properties
            if title:
                doc.core_properties.title = title
            if author:
                doc.core_properties.author = author

            # Ensure directory exists
            Path(file_path).parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(file_path)

            return {
                "success": True,
                "message": f"Document created at {file_path}",
                "file_path": file_path,
                "properties": {
                    "title": title,
                    "author": author,
                    "paragraphs": 0,
                    "runs": 0
                }
            }
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_text(file_path: str, text: str, paragraph_index: int | None = None, style: str | None = None) -> dict[str, Any]:
        """Add text to a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            if paragraph_index is None:
                # Add new paragraph at the end
                paragraph = doc.add_paragraph(text)
            else:
                # Insert at specific position
                if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                    return {"success": False, "error": f"Invalid paragraph index: {paragraph_index}"}

                # Insert new paragraph at specified index
                p = doc.paragraphs[paragraph_index]._element
                new_p = doc.add_paragraph(text)._element
                p.getparent().insert(p.getparent().index(p), new_p)
                paragraph = doc.paragraphs[paragraph_index]

            # Apply style if specified
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    logger.warning(f"Style '{style}' not found, using default")

            doc.save(file_path)

            return {
                "success": True,
                "message": f"Text added to document",
                "paragraph_index": len(doc.paragraphs) - 1 if paragraph_index is None else paragraph_index,
                "text": text
            }
        except Exception as e:
            logger.error(f"Error adding text: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_heading(file_path: str, text: str, level: int = 1) -> dict[str, Any]:
        """Add a heading to a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)
            heading = doc.add_heading(text, level)
            doc.save(file_path)

            return {
                "success": True,
                "message": f"Heading added to document",
                "text": text,
                "level": level,
                "paragraph_index": len(doc.paragraphs) - 1
            }
        except Exception as e:
            logger.error(f"Error adding heading: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def format_text(file_path: str, paragraph_index: int, run_index: int | None = None,
                   bold: bool | None = None, italic: bool | None = None, underline: bool | None = None,
                   font_size: int | None = None, font_name: str | None = None) -> dict[str, Any]:
        """Format text in a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"success": False, "error": f"Invalid paragraph index: {paragraph_index}"}

            paragraph = doc.paragraphs[paragraph_index]

            if run_index is None:
                # Format entire paragraph
                runs = paragraph.runs
            else:
                if run_index < 0 or run_index >= len(paragraph.runs):
                    return {"success": False, "error": f"Invalid run index: {run_index}"}
                runs = [paragraph.runs[run_index]]

            # Apply formatting
            for run in runs:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                if font_size is not None:
                    run.font.size = Pt(font_size)
                if font_name is not None:
                    run.font.name = font_name

            doc.save(file_path)

            return {
                "success": True,
                "message": f"Text formatted",
                "paragraph_index": paragraph_index,
                "run_index": run_index,
                "formatting_applied": {
                    "bold": bold,
                    "italic": italic,
                    "underline": underline,
                    "font_size": font_size,
                    "font_name": font_name
                }
            }
        except Exception as e:
            logger.error(f"Error formatting text: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_table(file_path: str, rows: int, cols: int, data: list[list[str]] | None = None,
                  headers: list[str] | None = None) -> dict[str, Any]:
        """Add a table to a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            # Add headers if provided
            if headers and len(headers) <= cols:
                for i, header in enumerate(headers):
                    table.cell(0, i).text = header
                    # Make header bold
                    for paragraph in table.cell(0, i).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

            # Add data if provided
            if data:
                start_row = 1 if headers else 0
                for row_idx, row_data in enumerate(data):
                    if row_idx + start_row >= rows:
                        break
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx >= cols:
                            break
                        table.cell(row_idx + start_row, col_idx).text = str(cell_data)

            doc.save(file_path)

            return {
                "success": True,
                "message": f"Table added to document",
                "rows": rows,
                "cols": cols,
                "has_headers": bool(headers),
                "has_data": bool(data)
            }
        except Exception as e:
            logger.error(f"Error adding table: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def analyze_document(file_path: str, include_structure: bool = True, include_formatting: bool = True,
                        include_statistics: bool = True) -> dict[str, Any]:
        """Analyze document content and structure."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)
            analysis = {"success": True}

            if include_structure:
                structure = {
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                    "headings": [],
                    "paragraphs_with_text": 0
                }

                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        structure["paragraphs_with_text"] += 1

                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        structure["headings"].append({
                            "index": i,
                            "text": para.text,
                            "level": para.style.name,
                            "style": para.style.name
                        })

                analysis["structure"] = structure

            if include_formatting:
                formatting = {
                    "styles_used": [],
                    "font_names": set(),
                    "font_sizes": set()
                }

                for para in doc.paragraphs:
                    if para.style.name not in formatting["styles_used"]:
                        formatting["styles_used"].append(para.style.name)

                    for run in para.runs:
                        if run.font.name:
                            formatting["font_names"].add(run.font.name)
                        if run.font.size:
                            formatting["font_sizes"].add(str(run.font.size))

                # Convert sets to lists for JSON serialization
                formatting["font_names"] = list(formatting["font_names"])
                formatting["font_sizes"] = list(formatting["font_sizes"])

                analysis["formatting"] = formatting

            if include_statistics:
                all_text = "\n".join([para.text for para in doc.paragraphs])
                words = all_text.split()

                statistics = {
                    "total_characters": len(all_text),
                    "total_words": len(words),
                    "total_sentences": len([s for s in all_text.split('.') if s.strip()]),
                    "average_words_per_paragraph": len(words) / max(len(doc.paragraphs), 1),
                    "longest_paragraph": max([len(para.text) for para in doc.paragraphs] + [0]),
                }

                analysis["statistics"] = statistics

            # Document properties
            analysis["properties"] = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None
            }

            return analysis
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def extract_text(file_path: str) -> dict[str, Any]:
        """Extract all text from a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            content = {
                "paragraphs": [],
                "tables": []
            }

            # Extract paragraph text
            for i, para in enumerate(doc.paragraphs):
                content["paragraphs"].append({
                    "index": i,
                    "text": para.text,
                    "style": para.style.name
                })

            # Extract table text
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                content["tables"].append({
                    "index": i,
                    "data": table_data,
                    "rows": len(table.rows),
                    "cols": len(table.columns) if table.rows else 0
                })

            return {
                "success": True,
                "content": content,
                "full_text": "\n".join([para.text for para in doc.paragraphs])
            }
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return {"success": False, "error": str(e)}


@server.list_tools()
async def handle_list_tools() -> list[Tool]:
    """List available DOCX tools."""
    return [
        Tool(
            name="create_document",
            description="Create a new DOCX document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path where the document will be saved"
                    },
                    "title": {
                        "type": "string",
                        "description": "Document title (optional)"
                    },
                    "author": {
                        "type": "string",
                        "description": "Document author (optional)"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="add_text",
            description="Add text to a document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the DOCX file"
                    },
                    "text": {
                        "type": "string",
                        "description": "Text to add"
                    },
                    "paragraph_index": {
                        "type": "integer",
                        "description": "Paragraph index to insert at (optional, defaults to end)"
                    },
                    "style": {
                        "type": "string",
                        "description": "Style to apply (optional)"
                    }
                },
                "required": ["file_path", "text"]
            }
        ),
        Tool(
            name="add_heading",
            description="Add a heading to a document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the DOCX file"
                    },
                    "text": {
                        "type": "string",
                        "description": "Heading text"
                    },
                    "level": {
                        "type": "integer",
                        "description": "Heading level (1-9)",
                        "minimum": 1,
                        "maximum": 9,
                        "default": 1
                    }
                },
                "required": ["file_path", "text"]
            }
        ),
        Tool(
            name="format_text",
            description="Format text in a document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the DOCX file"
                    },
                    "paragraph_index": {
                        "type": "integer",
                        "description": "Paragraph index to format"
                    },
                    "run_index": {
                        "type": "integer",
                        "description": "Run index within paragraph (optional, formats entire paragraph if not specified)"
                    },
                    "bold": {
                        "type": "boolean",
                        "description": "Make text bold (optional)"
                    },
                    "italic": {
                        "type": "boolean",
                        "description": "Make text italic (optional)"
                    },
                    "underline": {
                        "type": "boolean",
                        "description": "Underline text (optional)"
                    },
                    "font_size": {
                        "type": "integer",
                        "description": "Font size in points (optional)"
                    },
                    "font_name": {
                        "type": "string",
                        "description": "Font name (optional)"
                    }
                },
                "required": ["file_path", "paragraph_index"]
            }
        ),
        Tool(
            name="add_table",
            description="Add a table to a document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the DOCX file"
                    },
                    "rows": {
                        "type": "integer",
                        "description": "Number of rows",
                        "minimum": 1
                    },
                    "cols": {
                        "type": "integer",
                        "description": "Number of columns",
                        "minimum": 1
                    },
                    "data": {
                        "type": "array",
                        "items": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "description": "Table data (optional)"
                    },
                    "headers": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Column headers (optional)"
                    }
                },
                "required": ["file_path", "rows", "cols"]
            }
        ),
        Tool(
            name="analyze_document",
            description="Analyze document content, structure, and formatting",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the DOCX file"
                    },
                    "include_structure": {
                        "type": "boolean",
                        "description": "Include document structure analysis",
                        "default": True
                    },
                    "include_formatting": {
                        "type": "boolean",
                        "description": "Include formatting analysis",
                        "default": True
                    },
                    "include_statistics": {
                        "type": "boolean",
                        "description": "Include text statistics",
                        "default": True
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="extract_text",
            description="Extract all text content from a document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the DOCX file"
                    }
                },
                "required": ["file_path"]
            }
        )
    ]


@server.call_tool()
async def handle_call_tool(name: str, arguments: dict[str, Any]) -> Sequence[TextContent | ImageContent | EmbeddedResource]:
    """Handle tool calls."""
    try:
        doc_ops = DocumentOperation()

        if name == "create_document":
            request = CreateDocumentRequest(**arguments)
            result = doc_ops.create_document(
                file_path=request.file_path,
                title=request.title,
                author=request.author
            )

        elif name == "add_text":
            request = AddTextRequest(**arguments)
            result = doc_ops.add_text(
                file_path=request.file_path,
                text=request.text,
                paragraph_index=request.paragraph_index,
                style=request.style
            )

        elif name == "add_heading":
            request = AddHeadingRequest(**arguments)
            result = doc_ops.add_heading(
                file_path=request.file_path,
                text=request.text,
                level=request.level
            )

        elif name == "format_text":
            request = FormatTextRequest(**arguments)
            result = doc_ops.format_text(
                file_path=request.file_path,
                paragraph_index=request.paragraph_index,
                run_index=request.run_index,
                bold=request.bold,
                italic=request.italic,
                underline=request.underline,
                font_size=request.font_size,
                font_name=request.font_name
            )

        elif name == "add_table":
            request = AddTableRequest(**arguments)
            result = doc_ops.add_table(
                file_path=request.file_path,
                rows=request.rows,
                cols=request.cols,
                data=request.data,
                headers=request.headers
            )

        elif name == "analyze_document":
            request = AnalyzeDocumentRequest(**arguments)
            result = doc_ops.analyze_document(
                file_path=request.file_path,
                include_structure=request.include_structure,
                include_formatting=request.include_formatting,
                include_statistics=request.include_statistics
            )

        elif name == "extract_text":
            request = DocumentRequest(**arguments)
            result = doc_ops.extract_text(file_path=request.file_path)

        else:
            result = {"success": False, "error": f"Unknown tool: {name}"}

    except Exception as e:
        logger.error(f"Error in {name}: {str(e)}")
        result = {"success": False, "error": str(e)}

    return [TextContent(type="text", text=json.dumps(result, indent=2))]


async def main():
    """Main server entry point."""
    logger.info("Starting DOCX MCP Server...")

    from mcp.server.stdio import stdio_server

    logger.info("Waiting for MCP client connection...")
    async with stdio_server() as (read_stream, write_stream):
        logger.info("MCP client connected, starting server...")
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="docx-server",
                server_version="0.1.0",
                capabilities={
                    "tools": {},
                    "logging": {},
                },
            ),
        )


if __name__ == "__main__":
    asyncio.run(main())
