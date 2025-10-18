#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Location: ./mcp-servers/python/docx_server/src/docx_server/server_fastmcp.py
Copyright 2025
SPDX-License-Identifier: Apache-2.0
Authors: Mihai Criveti

DOCX MCP Server - FastMCP Implementation

A comprehensive MCP server for creating, editing, and analyzing Microsoft Word (.docx) documents.
Provides tools for document creation, text manipulation, formatting, and document analysis.
"""

import logging
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from fastmcp import FastMCP
from pydantic import Field

# Configure logging to stderr to avoid MCP protocol interference
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stderr)],
)
logger = logging.getLogger(__name__)

# Create FastMCP server instance
mcp = FastMCP("docx-server")


class DocumentOperation:
    """Handles document operations."""

    @staticmethod
    def create_document(
        file_path: str, title: str | None = None, author: str | None = None
    ) -> dict[str, Any]:
        """Create a new DOCX document."""
        try:
            # Create document
            doc = Document()

            # Set document properties
            if title:
                doc.core_properties.title = title
            if author:
                doc.core_properties.author = author

            # Ensure directory exists
            Path(file_path).parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(file_path)

            return {
                "success": True,
                "message": f"Document created at {file_path}",
                "file_path": file_path,
                "properties": {"title": title, "author": author, "paragraphs": 0, "runs": 0},
            }
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_text(
        file_path: str, text: str, paragraph_index: int | None = None, style: str | None = None
    ) -> dict[str, Any]:
        """Add text to a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            if paragraph_index is None:
                # Add new paragraph at the end
                paragraph = doc.add_paragraph(text)
            else:
                # Insert at specific position
                if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                    return {
                        "success": False,
                        "error": f"Invalid paragraph index: {paragraph_index}",
                    }

                # Insert new paragraph at specified index
                p = doc.paragraphs[paragraph_index]._element
                new_p = doc.add_paragraph(text)._element
                p.getparent().insert(p.getparent().index(p), new_p)
                paragraph = doc.paragraphs[paragraph_index]

            # Apply style if specified
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    logger.warning(f"Style '{style}' not found, using default")

            doc.save(file_path)

            return {
                "success": True,
                "message": "Text added to document",
                "paragraph_index": len(doc.paragraphs) - 1
                if paragraph_index is None
                else paragraph_index,
                "text": text,
            }
        except Exception as e:
            logger.error(f"Error adding text: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_heading(file_path: str, text: str, level: int = 1) -> dict[str, Any]:
        """Add a heading to a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)
            heading = doc.add_heading(text, level)
            doc.save(file_path)

            return {
                "success": True,
                "message": "Heading added to document",
                "text": text,
                "level": level,
                "paragraph_index": len(doc.paragraphs) - 1,
            }
        except Exception as e:
            logger.error(f"Error adding heading: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def format_text(
        file_path: str,
        paragraph_index: int,
        run_index: int | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
        underline: bool | None = None,
        font_size: int | None = None,
        font_name: str | None = None,
    ) -> dict[str, Any]:
        """Format text in a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"success": False, "error": f"Invalid paragraph index: {paragraph_index}"}

            paragraph = doc.paragraphs[paragraph_index]

            if run_index is None:
                # Format entire paragraph
                runs = paragraph.runs
            else:
                if run_index < 0 or run_index >= len(paragraph.runs):
                    return {"success": False, "error": f"Invalid run index: {run_index}"}
                runs = [paragraph.runs[run_index]]

            # Apply formatting
            for run in runs:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                if font_size is not None:
                    run.font.size = Pt(font_size)
                if font_name is not None:
                    run.font.name = font_name

            doc.save(file_path)

            return {
                "success": True,
                "message": "Text formatted",
                "paragraph_index": paragraph_index,
                "run_index": run_index,
                "formatting_applied": {
                    "bold": bold,
                    "italic": italic,
                    "underline": underline,
                    "font_size": font_size,
                    "font_name": font_name,
                },
            }
        except Exception as e:
            logger.error(f"Error formatting text: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def add_table(
        file_path: str,
        rows: int,
        cols: int,
        data: list[list[str]] | None = None,
        headers: list[str] | None = None,
    ) -> dict[str, Any]:
        """Add a table to a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"

            # Add headers if provided
            if headers and len(headers) <= cols:
                for i, header in enumerate(headers):
                    table.cell(0, i).text = header
                    # Make header bold
                    for paragraph in table.cell(0, i).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

            # Add data if provided
            if data:
                start_row = 1 if headers else 0
                for row_idx, row_data in enumerate(data):
                    if row_idx + start_row >= rows:
                        break
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx >= cols:
                            break
                        table.cell(row_idx + start_row, col_idx).text = str(cell_data)

            doc.save(file_path)

            return {
                "success": True,
                "message": "Table added to document",
                "rows": rows,
                "cols": cols,
                "has_headers": bool(headers),
                "has_data": bool(data),
            }
        except Exception as e:
            logger.error(f"Error adding table: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def analyze_document(
        file_path: str,
        include_structure: bool = True,
        include_formatting: bool = True,
        include_statistics: bool = True,
    ) -> dict[str, Any]:
        """Analyze document content and structure."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)
            analysis = {"success": True}

            if include_structure:
                structure = {
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                    "headings": [],
                    "paragraphs_with_text": 0,
                }

                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        structure["paragraphs_with_text"] += 1

                    # Check if it's a heading
                    if para.style.name.startswith("Heading"):
                        structure["headings"].append(
                            {
                                "index": i,
                                "text": para.text,
                                "level": para.style.name,
                                "style": para.style.name,
                            }
                        )

                analysis["structure"] = structure

            if include_formatting:
                formatting = {"styles_used": [], "font_names": set(), "font_sizes": set()}

                for para in doc.paragraphs:
                    if para.style.name not in formatting["styles_used"]:
                        formatting["styles_used"].append(para.style.name)

                    for run in para.runs:
                        if run.font.name:
                            formatting["font_names"].add(run.font.name)
                        if run.font.size:
                            formatting["font_sizes"].add(str(run.font.size))

                # Convert sets to lists for JSON serialization
                formatting["font_names"] = list(formatting["font_names"])
                formatting["font_sizes"] = list(formatting["font_sizes"])

                analysis["formatting"] = formatting

            if include_statistics:
                all_text = "\n".join([para.text for para in doc.paragraphs])
                words = all_text.split()

                statistics = {
                    "total_characters": len(all_text),
                    "total_words": len(words),
                    "total_sentences": len([s for s in all_text.split(".") if s.strip()]),
                    "average_words_per_paragraph": len(words) / max(len(doc.paragraphs), 1),
                    "longest_paragraph": max([len(para.text) for para in doc.paragraphs] + [0]),
                }

                analysis["statistics"] = statistics

            # Document properties
            analysis["properties"] = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created)
                if doc.core_properties.created
                else None,
                "modified": str(doc.core_properties.modified)
                if doc.core_properties.modified
                else None,
            }

            return analysis
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            return {"success": False, "error": str(e)}

    @staticmethod
    def extract_text(file_path: str) -> dict[str, Any]:
        """Extract all text from a document."""
        try:
            if not Path(file_path).exists():
                return {"success": False, "error": f"Document not found: {file_path}"}

            doc = Document(file_path)

            content = {"paragraphs": [], "tables": []}

            # Extract paragraph text
            for i, para in enumerate(doc.paragraphs):
                content["paragraphs"].append(
                    {"index": i, "text": para.text, "style": para.style.name}
                )

            # Extract table text
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text)
                    table_content.append(row_content)

                content["tables"].append(
                    {
                        "index": table_idx,
                        "content": table_content,
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                    }
                )

            return {
                "success": True,
                "content": content,
                "total_paragraphs": len(content["paragraphs"]),
                "total_tables": len(content["tables"]),
            }
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return {"success": False, "error": str(e)}


# Initialize the document operations handler
doc_ops = DocumentOperation()


@mcp.tool(description="Create a new DOCX document")
async def create_document(
    file_path: str = Field(..., description="Path where the document will be saved"),
    title: str | None = Field(None, description="Document title"),
    author: str | None = Field(None, description="Document author"),
) -> dict[str, Any]:
    """Create a new DOCX document with optional metadata."""
    return doc_ops.create_document(file_path, title, author)


@mcp.tool(description="Add text to a document")
async def add_text(
    file_path: str = Field(..., description="Path to the DOCX file"),
    text: str = Field(..., description="Text to add"),
    paragraph_index: int | None = Field(
        None, description="Paragraph index to insert at (None for end)"
    ),
    style: str | None = Field(None, description="Style to apply"),
) -> dict[str, Any]:
    """Add text to an existing DOCX document."""
    return doc_ops.add_text(file_path, text, paragraph_index, style)


@mcp.tool(description="Add a heading to a document")
async def add_heading(
    file_path: str = Field(..., description="Path to the DOCX file"),
    text: str = Field(..., description="Heading text"),
    level: int = Field(1, description="Heading level (1-9)", ge=1, le=9),
) -> dict[str, Any]:
    """Add a formatted heading to a document."""
    return doc_ops.add_heading(file_path, text, level)


@mcp.tool(description="Format text in a document")
async def format_text(
    file_path: str = Field(..., description="Path to the DOCX file"),
    paragraph_index: int = Field(..., description="Paragraph index to format"),
    run_index: int | None = Field(
        None, description="Run index within paragraph (None for entire paragraph)"
    ),
    bold: bool | None = Field(None, description="Make text bold"),
    italic: bool | None = Field(None, description="Make text italic"),
    underline: bool | None = Field(None, description="Underline text"),
    font_size: int | None = Field(None, description="Font size in points"),
    font_name: str | None = Field(None, description="Font name"),
) -> dict[str, Any]:
    """Apply formatting to text in a document."""
    return doc_ops.format_text(
        file_path, paragraph_index, run_index, bold, italic, underline, font_size, font_name
    )


@mcp.tool(description="Add a table to a document")
async def add_table(
    file_path: str = Field(..., description="Path to the DOCX file"),
    rows: int = Field(..., description="Number of rows", ge=1),
    cols: int = Field(..., description="Number of columns", ge=1),
    data: list[list[str]] | None = Field(None, description="Table data (optional)"),
    headers: list[str] | None = Field(None, description="Column headers (optional)"),
) -> dict[str, Any]:
    """Add a table to a document with optional data and headers."""
    return doc_ops.add_table(file_path, rows, cols, data, headers)


@mcp.tool(description="Analyze document structure and content")
async def analyze_document(
    file_path: str = Field(..., description="Path to the DOCX file"),
    include_structure: bool = Field(True, description="Include document structure analysis"),
    include_formatting: bool = Field(True, description="Include formatting analysis"),
    include_statistics: bool = Field(True, description="Include text statistics"),
) -> dict[str, Any]:
    """Analyze a document's structure, formatting, and statistics."""
    return doc_ops.analyze_document(
        file_path, include_structure, include_formatting, include_statistics
    )


@mcp.tool(description="Extract all text content from a document")
async def extract_text(
    file_path: str = Field(..., description="Path to the DOCX file"),
) -> dict[str, Any]:
    """Extract all text content from a DOCX document."""
    return doc_ops.extract_text(file_path)


def main():
    """Main entry point for the FastMCP server."""
    import argparse

    parser = argparse.ArgumentParser(description="DOCX FastMCP Server")
    parser.add_argument(
        "--transport",
        choices=["stdio", "http"],
        default="stdio",
        help="Transport mode (stdio or http)",
    )
    parser.add_argument("--host", default="0.0.0.0", help="HTTP host")
    parser.add_argument("--port", type=int, default=9004, help="HTTP port")

    args = parser.parse_args()

    if args.transport == "http":
        logger.info(f"Starting DOCX FastMCP Server on HTTP at {args.host}:{args.port}")
        mcp.run(transport="http", host=args.host, port=args.port)
    else:
        logger.info("Starting DOCX FastMCP Server on stdio")
        mcp.run()


if __name__ == "__main__":
    main()
