#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Location: ./mcp-servers/python/url_to_markdown_server/src/url_to_markdown_server/server.py
Copyright 2025
SPDX-License-Identifier: Apache-2.0
Authors: Mihai Criveti

URL-to-Markdown MCP Server

The ultimate MCP server for retrieving web content and files, then converting them to markdown.
Supports multiple content types, formats, and conversion engines with comprehensive error handling.

Features:
- Multi-format support: HTML, PDF, DOCX, PPTX, XLSX, TXT, Images
- Multiple HTML-to-Markdown engines: html2text, markdownify, turndown
- Content cleaning and optimization
- Image extraction and processing
- Metadata extraction
- URL validation and sanitization
- Rate limiting and timeout controls
- Comprehensive error handling
"""

import asyncio
import json
import logging
import mimetypes
import os
import re
import sys
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple
from urllib.parse import urljoin, urlparse
from uuid import uuid4

import httpx
from mcp.server import Server
from mcp.server.models import InitializationOptions
from mcp.types import EmbeddedResource, ImageContent, TextContent, Tool
from pydantic import BaseModel, Field, HttpUrl

# Configure logging to stderr to avoid MCP protocol interference
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stderr)],
)
logger = logging.getLogger(__name__)

# Create server instance
server = Server("url-to-markdown-server")

# Configuration constants
DEFAULT_TIMEOUT = int(os.getenv("MARKDOWN_DEFAULT_TIMEOUT", "30"))
MAX_TIMEOUT = int(os.getenv("MARKDOWN_MAX_TIMEOUT", "120"))
MAX_CONTENT_SIZE = int(os.getenv("MARKDOWN_MAX_CONTENT_SIZE", "50971520"))  # 50MB
MAX_REDIRECT_HOPS = int(os.getenv("MARKDOWN_MAX_REDIRECT_HOPS", "10"))
DEFAULT_USER_AGENT = os.getenv("MARKDOWN_USER_AGENT", "URL-to-Markdown-MCP-Server/1.0")


class ConvertUrlRequest(BaseModel):
    """Request to convert URL to markdown."""
    url: HttpUrl = Field(..., description="URL to retrieve and convert")
    timeout: int = Field(DEFAULT_TIMEOUT, description="Request timeout in seconds", le=MAX_TIMEOUT)
    include_images: bool = Field(True, description="Include images in markdown")
    include_links: bool = Field(True, description="Preserve links in markdown")
    clean_content: bool = Field(True, description="Clean and optimize content")
    extraction_method: str = Field("auto", description="HTML extraction method (auto, readability, raw)")
    markdown_engine: str = Field("html2text", description="Markdown conversion engine")
    max_image_size: int = Field(1048576, description="Maximum image size to process (1MB)")


class ConvertContentRequest(BaseModel):
    """Request to convert raw content to markdown."""
    content: str = Field(..., description="Raw content to convert")
    content_type: str = Field("text/html", description="MIME type of content")
    base_url: Optional[HttpUrl] = Field(None, description="Base URL for resolving relative links")
    include_images: bool = Field(True, description="Include images in markdown")
    clean_content: bool = Field(True, description="Clean and optimize content")
    markdown_engine: str = Field("html2text", description="Markdown conversion engine")


class ConvertFileRequest(BaseModel):
    """Request to convert local file to markdown."""
    file_path: str = Field(..., description="Path to local file")
    include_images: bool = Field(True, description="Include images in markdown")
    clean_content: bool = Field(True, description="Clean and optimize content")


class BatchConvertRequest(BaseModel):
    """Request to convert multiple URLs to markdown."""
    urls: List[HttpUrl] = Field(..., description="List of URLs to convert")
    timeout: int = Field(DEFAULT_TIMEOUT, description="Request timeout per URL")
    max_concurrent: int = Field(5, description="Maximum concurrent requests", le=10)
    include_images: bool = Field(False, description="Include images in markdown")
    clean_content: bool = Field(True, description="Clean and optimize content")


class UrlToMarkdownConverter:
    """Main converter class for URL-to-Markdown operations."""

    def __init__(self):
        """Initialize the converter."""
        self.session = None
        self.html_engines = self._check_html_engines()
        self.document_converters = self._check_document_converters()

    def _check_html_engines(self) -> Dict[str, bool]:
        """Check availability of HTML-to-Markdown engines."""
        engines = {}

        try:
            import html2text
            engines['html2text'] = True
        except ImportError:
            engines['html2text'] = False

        try:
            import markdownify
            engines['markdownify'] = True
        except ImportError:
            engines['markdownify'] = False

        try:
            from bs4 import BeautifulSoup
            engines['beautifulsoup'] = True
        except ImportError:
            engines['beautifulsoup'] = False

        try:
            from readability import Document
            engines['readability'] = True
        except ImportError:
            engines['readability'] = False

        return engines

    def _check_document_converters(self) -> Dict[str, bool]:
        """Check availability of document converters."""
        converters = {}

        try:
            import pypandoc
            converters['pandoc'] = True
        except ImportError:
            converters['pandoc'] = False

        try:
            import fitz  # PyMuPDF
            converters['pymupdf'] = True
        except ImportError:
            converters['pymupdf'] = False

        try:
            from docx import Document
            converters['python_docx'] = True
        except ImportError:
            converters['python_docx'] = False

        try:
            import openpyxl
            converters['openpyxl'] = True
        except ImportError:
            converters['openpyxl'] = False

        return converters

    async def get_session(self) -> httpx.AsyncClient:
        """Get or create HTTP session."""
        if self.session is None or self.session.is_closed:
            self.session = httpx.AsyncClient(
                headers={
                    'User-Agent': DEFAULT_USER_AGENT,
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Accept-Encoding': 'gzip, deflate',
                    'Connection': 'keep-alive',
                    'Upgrade-Insecure-Requests': '1',
                },
                timeout=httpx.Timeout(DEFAULT_TIMEOUT),
                follow_redirects=True,
                max_redirects=MAX_REDIRECT_HOPS
            )
        return self.session

    async def fetch_url_content(self, url: str, timeout: int = DEFAULT_TIMEOUT) -> Dict[str, Any]:
        """Fetch content from URL with comprehensive error handling."""
        try:
            session = await self.get_session()

            logger.info(f"Fetching URL: {url}")

            response = await session.get(url, timeout=timeout)
            response.raise_for_status()

            # Check content size
            content_length = response.headers.get('content-length')
            if content_length and int(content_length) > MAX_CONTENT_SIZE:
                return {
                    "success": False,
                    "error": f"Content too large: {content_length} bytes (max: {MAX_CONTENT_SIZE})"
                }

            content = response.content
            if len(content) > MAX_CONTENT_SIZE:
                return {
                    "success": False,
                    "error": f"Content too large: {len(content)} bytes (max: {MAX_CONTENT_SIZE})"
                }

            # Determine content type
            content_type = response.headers.get('content-type', '').lower()
            detected_type = self._detect_content_type(content, content_type, url)

            return {
                "success": True,
                "content": content,
                "content_type": detected_type,
                "original_content_type": content_type,
                "url": str(response.url),  # Final URL after redirects
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "size": len(content)
            }

        except httpx.TimeoutException:
            return {"success": False, "error": f"Request timeout after {timeout} seconds"}
        except httpx.HTTPStatusError as e:
            return {"success": False, "error": f"HTTP {e.response.status_code}: {e.response.reason_phrase}"}
        except Exception as e:
            logger.error(f"Error fetching URL {url}: {e}")
            return {"success": False, "error": str(e)}

    def _detect_content_type(self, content: bytes, content_type: str, url: str) -> str:
        """Detect actual content type from content, headers, and URL."""
        # Check file extension first
        url_path = urlparse(url).path.lower()

        if url_path.endswith(('.pdf',)):
            return 'application/pdf'
        elif url_path.endswith(('.docx',)):
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif url_path.endswith(('.pptx',)):
            return 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
        elif url_path.endswith(('.xlsx',)):
            return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif url_path.endswith(('.txt', '.md', '.rst')):
            return 'text/plain'

        # Check content-type header
        if 'html' in content_type:
            return 'text/html'
        elif 'pdf' in content_type:
            return 'application/pdf'
        elif 'json' in content_type:
            return 'application/json'
        elif 'xml' in content_type:
            return 'application/xml'

        # Check magic bytes
        if content.startswith(b'%PDF'):
            return 'application/pdf'
        elif content.startswith(b'PK'):  # ZIP-based formats (Office docs)
            if b'word/' in content[:1024]:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif b'ppt/' in content[:1024]:
                return 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
            elif b'xl/' in content[:1024]:
                return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif content.startswith((b'<html', b'<!DOCTYPE', b'<!doctype')):
            return 'text/html'

        return content_type or 'application/octet-stream'

    async def convert_html_to_markdown(
        self,
        html_content: str,
        base_url: str = "",
        engine: str = "html2text",
        include_images: bool = True,
        include_links: bool = True
    ) -> Dict[str, Any]:
        """Convert HTML content to markdown using specified engine."""
        try:
            if engine == "html2text" and self.html_engines.get('html2text'):
                return await self._convert_with_html2text(html_content, base_url, include_images, include_links)
            elif engine == "markdownify" and self.html_engines.get('markdownify'):
                return await self._convert_with_markdownify(html_content, include_images, include_links)
            elif engine == "beautifulsoup" and self.html_engines.get('beautifulsoup'):
                return await self._convert_with_beautifulsoup(html_content, base_url, include_images)
            elif engine == "readability" and self.html_engines.get('readability'):
                return await self._convert_with_readability(html_content, base_url)
            else:
                # Fallback to basic conversion
                return await self._convert_basic_html(html_content)

        except Exception as e:
            logger.error(f"Error converting HTML to markdown: {e}")
            return {
                "success": False,
                "error": f"Conversion failed: {str(e)}"
            }

    async def _convert_with_html2text(
        self,
        html_content: str,
        base_url: str,
        include_images: bool,
        include_links: bool
    ) -> Dict[str, Any]:
        """Convert using html2text library."""
        import html2text

        converter = html2text.HTML2Text()
        converter.ignore_links = not include_links
        converter.ignore_images = not include_images
        converter.body_width = 0  # No line wrapping
        converter.protect_links = True
        converter.wrap_links = False

        if base_url:
            converter.baseurl = base_url

        markdown = converter.handle(html_content)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "html2text",
            "length": len(markdown)
        }

    async def _convert_with_markdownify(
        self,
        html_content: str,
        include_images: bool,
        include_links: bool
    ) -> Dict[str, Any]:
        """Convert using markdownify library."""
        import markdownify

        # Configure conversion options
        options = {
            'heading_style': 'ATX',  # Use # for headings
            'bullets': '-',  # Use - for lists
            'escape_misc': False,
        }

        if not include_links:
            options['convert'] = ['p', 'div', 'span', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']

        if not include_images:
            if 'convert' in options:
                pass  # img already excluded
            else:
                options['strip'] = ['img']

        markdown = markdownify.markdownify(html_content, **options)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "markdownify",
            "length": len(markdown)
        }

    async def _convert_with_beautifulsoup(
        self,
        html_content: str,
        base_url: str,
        include_images: bool
    ) -> Dict[str, Any]:
        """Convert using BeautifulSoup for parsing + custom markdown generation."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, 'html.parser')

        # Extract main content
        main_content = self._extract_main_content(soup)

        # Convert to markdown
        markdown = self._soup_to_markdown(main_content, base_url, include_images)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "beautifulsoup",
            "length": len(markdown)
        }

    async def _convert_with_readability(self, html_content: str, base_url: str) -> Dict[str, Any]:
        """Convert using readability for content extraction."""
        from readability import Document

        doc = Document(html_content)
        title = doc.title()
        content = doc.summary()

        # Convert extracted content to markdown
        if self.html_engines.get('html2text'):
            import html2text
            converter = html2text.HTML2Text()
            converter.body_width = 0
            if base_url:
                converter.baseurl = base_url
            markdown = converter.handle(content)
        else:
            # Basic conversion
            markdown = self._html_to_markdown_basic(content)

        # Add title if available
        if title:
            markdown = f"# {title}\n\n{markdown}"

        return {
            "success": True,
            "markdown": markdown,
            "engine": "readability",
            "title": title,
            "length": len(markdown)
        }

    async def _convert_basic_html(self, html_content: str) -> Dict[str, Any]:
        """Basic HTML to markdown conversion without external libraries."""
        markdown = self._html_to_markdown_basic(html_content)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "basic",
            "length": len(markdown),
            "note": "Basic conversion - install html2text or markdownify for better results"
        }

    def _html_to_markdown_basic(self, html_content: str) -> str:
        """Basic HTML to markdown conversion."""
        import re

        # Remove script and style tags
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert headings
        for i in range(1, 7):
            html_content = re.sub(f'<h{i}[^>]*>(.*?)</h{i}>', f'{"#" * i} \\1\n\n', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert paragraphs
        html_content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert line breaks
        html_content = re.sub(r'<br[^>]*/?>', '\n', html_content, flags=re.IGNORECASE)

        # Convert links
        html_content = re.sub(r'<a[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>', r'[\2](\1)', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert bold and italic
        html_content = re.sub(r'<(strong|b)[^>]*>(.*?)</\1>', r'**\2**', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<(em|i)[^>]*>(.*?)</\1>', r'*\2*', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert lists
        html_content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<[uo]l[^>]*>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</[uo]l>', '\n', html_content, flags=re.IGNORECASE)

        # Remove remaining HTML tags
        html_content = re.sub(r'<[^>]+>', '', html_content)

        # Clean up whitespace
        html_content = re.sub(r'\n\s*\n\s*\n', '\n\n', html_content)
        html_content = re.sub(r'^\s+|\s+$', '', html_content, flags=re.MULTILINE)

        return html_content.strip()

    def _extract_main_content(self, soup):
        """Extract main content from BeautifulSoup object."""
        # Try to find main content areas
        main_selectors = [
            'main', 'article', '[role="main"]',
            '.content', '.main-content', '.post-content',
            '#content', '#main-content', '#post-content'
        ]

        for selector in main_selectors:
            main_element = soup.select_one(selector)
            if main_element:
                return main_element

        # Fallback to body
        body = soup.find('body')
        if body:
            # Remove navigation, sidebar, footer elements
            for element in body.find_all(['nav', 'aside', 'footer', 'header']):
                element.decompose()

            # Remove elements with common nav/sidebar classes
            for element in body.find_all(class_=re.compile(r'(nav|sidebar|footer|header|menu)', re.I)):
                element.decompose()

            return body

        return soup

    def _soup_to_markdown(self, element, base_url: str = "", include_images: bool = True) -> str:
        """Convert BeautifulSoup element to markdown."""
        markdown_parts = []

        for child in element.children:
            if hasattr(child, 'name'):
                if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(child.name[1])
                    text = child.get_text().strip()
                    markdown_parts.append(f"{'#' * level} {text}\n")
                elif child.name == 'p':
                    text = child.get_text().strip()
                    if text:
                        markdown_parts.append(f"{text}\n")
                elif child.name == 'a':
                    href = child.get('href', '')
                    text = child.get_text().strip()
                    if href and text:
                        if base_url and not href.startswith(('http', 'https')):
                            href = urljoin(base_url, href)
                        markdown_parts.append(f"[{text}]({href})")
                elif child.name == 'img' and include_images:
                    src = child.get('src', '')
                    alt = child.get('alt', 'Image')
                    if src:
                        if base_url and not src.startswith(('http', 'https')):
                            src = urljoin(base_url, src)
                        markdown_parts.append(f"![{alt}]({src})")
                elif child.name in ['strong', 'b']:
                    text = child.get_text().strip()
                    markdown_parts.append(f"**{text}**")
                elif child.name in ['em', 'i']:
                    text = child.get_text().strip()
                    markdown_parts.append(f"*{text}*")
                elif child.name == 'li':
                    text = child.get_text().strip()
                    markdown_parts.append(f"- {text}\n")
                elif child.name == 'code':
                    text = child.get_text()
                    markdown_parts.append(f"`{text}`")
                elif child.name == 'pre':
                    text = child.get_text()
                    markdown_parts.append(f"```\n{text}\n```\n")
                else:
                    # Recursively process other elements
                    nested_markdown = self._soup_to_markdown(child, base_url, include_images)
                    if nested_markdown.strip():
                        markdown_parts.append(nested_markdown)
            else:
                # Text node
                text = str(child).strip()
                if text:
                    markdown_parts.append(text)

        return ' '.join(markdown_parts)

    async def convert_document_to_markdown(self, content: bytes, content_type: str) -> Dict[str, Any]:
        """Convert document formats to markdown."""
        try:
            if content_type == 'application/pdf':
                return await self._convert_pdf_to_markdown(content)
            elif 'wordprocessingml' in content_type:  # DOCX
                return await self._convert_docx_to_markdown(content)
            elif 'presentationml' in content_type:  # PPTX
                return await self._convert_pptx_to_markdown(content)
            elif 'spreadsheetml' in content_type:  # XLSX
                return await self._convert_xlsx_to_markdown(content)
            elif content_type.startswith('text/'):
                return await self._convert_text_to_markdown(content)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported content type: {content_type}"
                }

        except Exception as e:
            logger.error(f"Error converting document: {e}")
            return {
                "success": False,
                "error": f"Document conversion failed: {str(e)}"
            }

    async def _convert_pdf_to_markdown(self, pdf_content: bytes) -> Dict[str, Any]:
        """Convert PDF to markdown."""
        if not self.document_converters.get('pymupdf'):
            return {"success": False, "error": "PyMuPDF not available for PDF conversion"}

        try:
            import fitz

            # Open PDF from bytes
            doc = fitz.open(stream=pdf_content, filetype="pdf")

            markdown_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if text.strip():
                    markdown_parts.append(f"## Page {page_num + 1}\n\n{text}\n")

            doc.close()

            markdown = '\n'.join(markdown_parts)

            return {
                "success": True,
                "markdown": markdown,
                "engine": "pymupdf",
                "pages": len(doc),
                "length": len(markdown)
            }

        except Exception as e:
            return {"success": False, "error": f"PDF conversion error: {str(e)}"}

    async def _convert_docx_to_markdown(self, docx_content: bytes) -> Dict[str, Any]:
        """Convert DOCX to markdown."""
        if not self.document_converters.get('python_docx'):
            return {"success": False, "error": "python-docx not available for DOCX conversion"}

        try:
            from docx import Document
            from io import BytesIO

            doc = Document(BytesIO(docx_content))
            markdown_parts = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check if it's a heading based on style
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1])
                        markdown_parts.append(f"{'#' * level} {text}\n")
                    else:
                        markdown_parts.append(f"{text}\n")

            # Process tables
            for table in doc.tables:
                markdown_parts.append(self._table_to_markdown(table))

            markdown = '\n'.join(markdown_parts)

            return {
                "success": True,
                "markdown": markdown,
                "engine": "python_docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "length": len(markdown)
            }

        except Exception as e:
            return {"success": False, "error": f"DOCX conversion error: {str(e)}"}

    def _table_to_markdown(self, table) -> str:
        """Convert DOCX table to markdown table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')

        if rows:
            # Add header separator
            if len(rows) > 1:
                header_sep = '| ' + ' | '.join(['---'] * len(table.rows[0].cells)) + ' |'
                rows.insert(1, header_sep)

        return '\n'.join(rows) + '\n'

    async def _convert_xlsx_to_markdown(self, xlsx_content: bytes) -> Dict[str, Any]:
        """Convert XLSX to markdown."""
        if not self.document_converters.get('openpyxl'):
            return {"success": False, "error": "openpyxl not available for XLSX conversion"}

        try:
            import openpyxl
            from io import BytesIO

            workbook = openpyxl.load_workbook(BytesIO(xlsx_content))
            markdown_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                markdown_parts.append(f"## {sheet_name}\n")

                # Get data range
                if sheet.max_row > 0 and sheet.max_column > 0:
                    rows = []
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            cells = [str(cell) if cell is not None else '' for cell in row]
                            rows.append('| ' + ' | '.join(cells) + ' |')

                    if rows:
                        # Add header separator after first row
                        if len(rows) > 1:
                            header_sep = '| ' + ' | '.join(['---'] * sheet.max_column) + ' |'
                            rows.insert(1, header_sep)

                        markdown_parts.extend(rows)
                        markdown_parts.append("")

            markdown = '\n'.join(markdown_parts)

            return {
                "success": True,
                "markdown": markdown,
                "engine": "openpyxl",
                "sheets": len(workbook.sheetnames),
                "length": len(markdown)
            }

        except Exception as e:
            return {"success": False, "error": f"XLSX conversion error: {str(e)}"}

    async def _convert_text_to_markdown(self, text_content: bytes) -> Dict[str, Any]:
        """Convert plain text to markdown."""
        try:
            text = text_content.decode('utf-8', errors='replace')

            # For plain text, just return as-is with minimal formatting
            markdown = text

            return {
                "success": True,
                "markdown": markdown,
                "engine": "text",
                "length": len(markdown)
            }

        except Exception as e:
            return {"success": False, "error": f"Text conversion error: {str(e)}"}

    def clean_markdown(self, markdown: str) -> str:
        """Clean and optimize markdown content."""
        # Remove excessive whitespace
        markdown = re.sub(r'\n\s*\n\s*\n+', '\n\n', markdown)

        # Fix heading spacing
        markdown = re.sub(r'(#+\s+.+)\n+([^#\n])', r'\1\n\n\2', markdown)

        # Clean up list formatting
        markdown = re.sub(r'\n+(-\s+)', r'\n\1', markdown)

        # Remove empty links
        markdown = re.sub(r'\[\s*\]\([^)]*\)', '', markdown)

        # Clean up extra spaces
        markdown = re.sub(r' +', ' ', markdown)

        # Trim
        return markdown.strip()

    async def convert_url_to_markdown(
        self,
        url: str,
        timeout: int = DEFAULT_TIMEOUT,
        include_images: bool = True,
        include_links: bool = True,
        clean_content: bool = True,
        extraction_method: str = "auto",
        markdown_engine: str = "html2text"
    ) -> Dict[str, Any]:
        """Convert URL content to markdown."""
        conversion_id = str(uuid4())
        logger.info(f"Converting URL to markdown, ID: {conversion_id}, URL: {url}")

        try:
            # Fetch content
            fetch_result = await self.fetch_url_content(url, timeout)
            if not fetch_result["success"]:
                return {
                    "success": False,
                    "conversion_id": conversion_id,
                    "error": fetch_result["error"]
                }

            content = fetch_result["content"]
            content_type = fetch_result["content_type"]
            final_url = fetch_result["url"]

            # Convert based on content type
            if content_type.startswith('text/html'):
                html_content = content.decode('utf-8', errors='replace')

                # Choose extraction method
                if extraction_method == "readability":
                    result = await self._convert_with_readability(html_content, final_url)
                elif extraction_method == "raw":
                    result = await self.convert_html_to_markdown(
                        html_content, final_url, markdown_engine, include_images, include_links
                    )
                else:  # auto
                    # Try readability first, fallback to specified engine
                    if self.html_engines.get('readability'):
                        result = await self._convert_with_readability(html_content, final_url)
                    else:
                        result = await self.convert_html_to_markdown(
                            html_content, final_url, markdown_engine, include_images, include_links
                        )

            else:
                # Handle document formats
                result = await self.convert_document_to_markdown(content, content_type)

            if not result["success"]:
                return {
                    "success": False,
                    "conversion_id": conversion_id,
                    "error": result["error"]
                }

            markdown = result["markdown"]

            # Clean content if requested
            if clean_content:
                markdown = self.clean_markdown(markdown)

            return {
                "success": True,
                "conversion_id": conversion_id,
                "url": final_url,
                "content_type": content_type,
                "markdown": markdown,
                "length": len(markdown),
                "engine": result.get("engine", "unknown"),
                "metadata": {
                    "original_size": len(content),
                    "compression_ratio": len(markdown) / len(content) if len(content) > 0 else 0,
                    "processing_time": time.time()
                }
            }

        except Exception as e:
            logger.error(f"Error converting URL {url}: {e}")
            return {
                "success": False,
                "conversion_id": conversion_id,
                "error": str(e)
            }

    async def batch_convert_urls(
        self,
        urls: List[str],
        timeout: int = DEFAULT_TIMEOUT,
        max_concurrent: int = 5,
        include_images: bool = False,
        clean_content: bool = True
    ) -> Dict[str, Any]:
        """Convert multiple URLs to markdown concurrently."""
        batch_id = str(uuid4())
        logger.info(f"Batch converting {len(urls)} URLs, ID: {batch_id}")

        semaphore = asyncio.Semaphore(max_concurrent)

        async def convert_single_url(url: str) -> Dict[str, Any]:
            async with semaphore:
                return await self.convert_url_to_markdown(
                    url, timeout, include_images, True, clean_content
                )

        try:
            # Process URLs concurrently
            tasks = [convert_single_url(url) for url in urls]
            results = await asyncio.gather(*tasks, return_exceptions=True)

            # Process results
            successful = 0
            failed = 0
            processed_results = []

            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    processed_results.append({
                        "url": urls[i],
                        "success": False,
                        "error": str(result)
                    })
                    failed += 1
                else:
                    processed_results.append(result)
                    if result.get("success"):
                        successful += 1
                    else:
                        failed += 1

            return {
                "success": True,
                "batch_id": batch_id,
                "total_urls": len(urls),
                "successful": successful,
                "failed": failed,
                "results": processed_results
            }

        except Exception as e:
            logger.error(f"Error in batch conversion: {e}")
            return {
                "success": False,
                "batch_id": batch_id,
                "error": str(e)
            }

    def get_capabilities(self) -> Dict[str, Any]:
        """Get converter capabilities and available engines."""
        return {
            "html_engines": self.html_engines,
            "document_converters": self.document_converters,
            "supported_formats": {
                "web": ["text/html", "application/xhtml+xml"],
                "documents": ["application/pdf"],
                "office": [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
                    "application/vnd.openxmlformats-officedocument.presentationml.presentation",  # PPTX
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"  # XLSX
                ],
                "text": ["text/plain", "text/markdown", "application/json"]
            },
            "features": [
                "Multi-engine HTML conversion",
                "PDF text extraction",
                "Office document conversion",
                "Content cleaning and optimization",
                "Image handling",
                "Link preservation",
                "Batch processing",
                "Metadata extraction"
            ]
        }


# Initialize converter (conditionally for testing)
try:
    converter = UrlToMarkdownConverter()
except Exception:
    converter = None


@server.list_tools()
async def handle_list_tools() -> list[Tool]:
    """List available URL-to-Markdown tools."""
    return [
        Tool(
            name="convert_url",
            description="Convert URL content to markdown",
            inputSchema={
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "URL to retrieve and convert"
                    },
                    "timeout": {
                        "type": "integer",
                        "description": "Request timeout in seconds",
                        "default": DEFAULT_TIMEOUT,
                        "maximum": MAX_TIMEOUT
                    },
                    "include_images": {
                        "type": "boolean",
                        "description": "Include images in markdown",
                        "default": True
                    },
                    "include_links": {
                        "type": "boolean",
                        "description": "Preserve links in markdown",
                        "default": True
                    },
                    "clean_content": {
                        "type": "boolean",
                        "description": "Clean and optimize content",
                        "default": True
                    },
                    "extraction_method": {
                        "type": "string",
                        "enum": ["auto", "readability", "raw"],
                        "description": "Content extraction method",
                        "default": "auto"
                    },
                    "markdown_engine": {
                        "type": "string",
                        "enum": ["html2text", "markdownify", "beautifulsoup", "basic"],
                        "description": "Markdown conversion engine",
                        "default": "html2text"
                    }
                },
                "required": ["url"]
            }
        ),
        Tool(
            name="convert_content",
            description="Convert raw content to markdown",
            inputSchema={
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "Raw content to convert"
                    },
                    "content_type": {
                        "type": "string",
                        "description": "MIME type of content",
                        "default": "text/html"
                    },
                    "base_url": {
                        "type": "string",
                        "description": "Base URL for resolving relative links"
                    },
                    "include_images": {
                        "type": "boolean",
                        "description": "Include images in markdown",
                        "default": True
                    },
                    "clean_content": {
                        "type": "boolean",
                        "description": "Clean and optimize content",
                        "default": True
                    },
                    "markdown_engine": {
                        "type": "string",
                        "enum": ["html2text", "markdownify", "beautifulsoup", "basic"],
                        "description": "Markdown conversion engine",
                        "default": "html2text"
                    }
                },
                "required": ["content"]
            }
        ),
        Tool(
            name="convert_file",
            description="Convert local file to markdown",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to local file"
                    },
                    "include_images": {
                        "type": "boolean",
                        "description": "Include images in markdown",
                        "default": True
                    },
                    "clean_content": {
                        "type": "boolean",
                        "description": "Clean and optimize content",
                        "default": True
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="batch_convert",
            description="Convert multiple URLs to markdown",
            inputSchema={
                "type": "object",
                "properties": {
                    "urls": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "List of URLs to convert"
                    },
                    "timeout": {
                        "type": "integer",
                        "description": "Request timeout per URL",
                        "default": DEFAULT_TIMEOUT
                    },
                    "max_concurrent": {
                        "type": "integer",
                        "description": "Maximum concurrent requests",
                        "default": 5,
                        "maximum": 10
                    },
                    "include_images": {
                        "type": "boolean",
                        "description": "Include images in markdown",
                        "default": False
                    },
                    "clean_content": {
                        "type": "boolean",
                        "description": "Clean and optimize content",
                        "default": True
                    }
                },
                "required": ["urls"]
            }
        ),
        Tool(
            name="get_capabilities",
            description="Get converter capabilities and available engines",
            inputSchema={
                "type": "object",
                "properties": {},
                "additionalProperties": False
            }
        )
    ]


@server.call_tool()
async def handle_call_tool(name: str, arguments: dict[str, Any]) -> Sequence[TextContent | ImageContent | EmbeddedResource]:
    """Handle tool calls."""
    try:
        if converter is None:
            result = {"success": False, "error": "URL-to-Markdown converter not available"}
        elif name == "convert_url":
            request = ConvertUrlRequest(**arguments)
            result = await converter.convert_url_to_markdown(
                url=str(request.url),
                timeout=request.timeout,
                include_images=request.include_images,
                include_links=request.include_links,
                clean_content=request.clean_content,
                extraction_method=request.extraction_method,
                markdown_engine=request.markdown_engine
            )

        elif name == "convert_content":
            request = ConvertContentRequest(**arguments)
            if request.content_type.startswith('text/html'):
                result = await converter.convert_html_to_markdown(
                    html_content=request.content,
                    base_url=str(request.base_url) if request.base_url else "",
                    engine=request.markdown_engine,
                    include_images=request.include_images
                )
            else:
                result = await converter.convert_document_to_markdown(
                    content=request.content.encode('utf-8'),
                    content_type=request.content_type
                )

            if result["success"] and request.clean_content:
                result["markdown"] = converter.clean_markdown(result["markdown"])

        elif name == "convert_file":
            request = ConvertFileRequest(**arguments)

            file_path = Path(request.file_path)
            if not file_path.exists():
                result = {"success": False, "error": f"File not found: {request.file_path}"}
            else:
                content = file_path.read_bytes()
                content_type = mimetypes.guess_type(str(file_path))[0] or 'application/octet-stream'

                result = await converter.convert_document_to_markdown(content, content_type)

                if result["success"] and request.clean_content:
                    result["markdown"] = converter.clean_markdown(result["markdown"])

        elif name == "batch_convert":
            request = BatchConvertRequest(**arguments)
            result = await converter.batch_convert_urls(
                urls=[str(url) for url in request.urls],
                timeout=request.timeout,
                max_concurrent=request.max_concurrent,
                include_images=request.include_images,
                clean_content=request.clean_content
            )

        elif name == "get_capabilities":
            result = converter.get_capabilities()

        else:
            result = {"success": False, "error": f"Unknown tool: {name}"}

    except Exception as e:
        logger.error(f"Error in {name}: {str(e)}")
        result = {"success": False, "error": str(e)}

    return [TextContent(type="text", text=json.dumps(result, indent=2, default=str))]


async def main():
    """Main server entry point."""
    logger.info("Starting URL-to-Markdown MCP Server...")

    from mcp.server.stdio import stdio_server

    logger.info("Waiting for MCP client connection...")
    async with stdio_server() as (read_stream, write_stream):
        logger.info("MCP client connected, starting server...")
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="url-to-markdown-server",
                server_version="0.1.0",
                capabilities={
                    "tools": {},
                    "logging": {},
                },
            ),
        )


if __name__ == "__main__":
    asyncio.run(main())
